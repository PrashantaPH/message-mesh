"""Generate a Word document documenting the MessageMesh Admin end-to-end flow."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(BASE, "admin-screenshots")
OUT = os.path.join(BASE, "MessageMesh-Admin-Flow-v1.docx")

ACCENT = RGBColor(0x1F, 0x6F, 0xEB)
MUTED = RGBColor(0x55, 0x5F, 0x6D)

steps = [
    {
        "img": "01-login.png",
        "title": "Admin Sign In",
        "url": "http://localhost:5173/login",
        "desc": "The admin journey begins at the standard sign-in screen. The administrator "
                "authenticates with the privileged account (username \"admin\"). Authentication "
                "is JWT-based; the returned token carries the ADMIN role claim, which unlocks the "
                "admin-only navigation and routes. All admin API calls and the WebSocket "
                "connection reuse this token for authorization.",
    },
    {
        "img": "02-admin-menu.png",
        "title": "Admin Navigation Menu",
        "url": "http://localhost:5173/chats",
        "desc": "After signing in, the admin lands on the chat workspace like any user, but the "
                "profile menu (opened from the avatar) reveals privileged entries. An \"Admin\" "
                "badge next to the display name confirms the elevated role, and a dedicated "
                "section links directly to the Admin dashboard, Manage users, Manage "
                "conversations, and the Audit log, alongside Profile & settings and Sign out. "
                "These options are rendered only when the current user has the ADMIN role.",
    },
    {
        "img": "03-dashboard.png",
        "title": "Admin Dashboard & Platform Analytics",
        "url": "http://localhost:5173/admin",
        "desc": "The Admin dashboard is the control center. Three shortcut cards summarize the "
                "platform and jump into each management area — registered users, conversations "
                "and recorded audit events, all fetched live from the backend. Below them, a "
                "\"Platform overview\" section renders twelve real-time metric tiles powered by "
                "the /api/admin/stats endpoint: active users (of total), online now, "
                "administrators, inactive users, new users in the last 7 days, group chats, "
                "direct chats, deleted chats, total messages, messages in the last 7 days, total "
                "conversations and audit events. The route is protected by a RequireAdmin guard, "
                "so non-admin users are redirected away.",
    },
    {
        "img": "04-users.png",
        "title": "Manage Users",
        "url": "http://localhost:5173/admin/users",
        "desc": "The Users page lists every registered account in a table showing a selection "
                "checkbox, avatar, display name, handle, role, active/inactive status, "
                "conversation count and join date. A debounced search box filters by name or "
                "username, and it is paired with a Role filter (All / Admins / Users), a Status "
                "filter (All / Active / Inactive) and a page-size selector (10 / 20 / 50 / 100 "
                "per page). The toolbar also offers \"Add user\" to create an account and "
                "\"Export CSV\" to download the current directory. Presence dots indicate who is "
                "currently online.",
    },
    {
        "img": "05-user-actions.png",
        "title": "User Actions",
        "url": "http://localhost:5173/admin/users",
        "desc": "Each user row exposes an actions menu with the full set of moderation controls: "
                "View details (opens the profile drawer), promote to Admin or demote to User "
                "(role management), activate or deactivate the account (access control), Revoke "
                "sessions to force the user to log out everywhere, Reset password, and "
                "permanently Delete user. Destructive and role-changing actions open a "
                "confirmation dialog, and guard rails prevent unsafe self-actions (e.g. demoting "
                "or deleting your own account). Every action shows a success/error toast.",
    },
    {
        "img": "06-create-user.png",
        "title": "Create a New User",
        "url": "http://localhost:5173/admin/users",
        "desc": "The \"Add user\" button opens a modal for provisioning an account directly from "
                "the admin console. The admin supplies a username, a display name, a temporary "
                "password and the initial role (User or Admin). On submit, the backend creates "
                "the account and the users table refreshes to include it. This lets "
                "administrators onboard people without requiring them to self-register.",
    },
    {
        "img": "07-user-detail.png",
        "title": "User Details & Force Logout",
        "url": "http://localhost:5173/admin/users",
        "desc": "Selecting \"View details\" slides out a drawer with a full profile: avatar, "
                "display name, handle, and role / status / presence badges, plus the join date "
                "and total conversation count. A \"Force logout (revoke sessions)\" button "
                "immediately invalidates the user's active sessions. The drawer also lists the "
                "conversations the user belongs to, each with its type and member / message "
                "counts, giving the admin quick context on the account's activity.",
    },
    {
        "img": "08-bulk-actions.png",
        "title": "Bulk User Actions",
        "url": "http://localhost:5173/admin/users",
        "desc": "The header checkbox selects every visible user and per-row checkboxes select "
                "individual accounts. As soon as one or more rows are selected, a bulk-action "
                "bar appears showing the selection count (\"3 selected\") with Activate, "
                "Deactivate and Delete operations that apply to the whole selection at once, plus "
                "a Clear button to reset it. This makes large-scale account maintenance far "
                "faster than acting on users one by one.",
    },
    {
        "img": "09-conversations.png",
        "title": "Manage Conversations",
        "url": "http://localhost:5173/admin/conversations",
        "desc": "The Conversations page provides oversight of every chat on the platform. The "
                "table lists the title, type (Group or Direct), member count, message count, "
                "creation date and status (Active/Deleted). Admins search by title and narrow "
                "results with a Type filter (Group / Direct) and a Status filter (Active / "
                "Deleted). Each row offers \"View messages\" for moderation and \"Delete "
                "conversation\"; conversations already soft-deleted expose a \"Restore\" action "
                "so a removed chat can be brought back.",
    },
    {
        "img": "10-conversation-messages.png",
        "title": "Message Moderation",
        "url": "http://localhost:5173/admin/conversations",
        "desc": "Choosing \"View messages\" opens a moderation drawer that streams the full "
                "message history of the selected conversation — each entry showing the sender, "
                "timestamp and content. From here an admin can delete an individual message "
                "rather than the entire conversation, enabling fine-grained moderation of "
                "specific inappropriate content while leaving the rest of the chat intact.",
    },
    {
        "img": "11-audit-log.png",
        "title": "Audit Log",
        "url": "http://localhost:5173/admin/audit",
        "desc": "The Audit log records administrative and membership activity across the platform "
                "for accountability and compliance. Each event captures when it happened, the "
                "actor, the action (e.g. MEMBER_ROLE_CHANGED, MEMBER_ADDED, CONVERSATION_RENAMED, "
                "MEMBER_LEFT), the target entity and additional details. Admins can filter by "
                "actor, by action and by a from/to date range to investigate specific activity, "
                "and export the filtered results to CSV for offline review or reporting.",
    },
    {
        "img": "12-dark-dashboard.png",
        "title": "Dark Mode",
        "url": "http://localhost:5173/admin",
        "desc": "The color-mode toggle in the admin header switches the entire admin area between "
                "light and dark themes. All admin surfaces — the dashboard cards and metric "
                "tiles, tables, filters, drawers, menus and dialogs — re-style consistently "
                "while preserving layout and readability, which is useful for long moderation "
                "sessions in low-light environments.",
    },
]

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

# ---- Title page ----
title = doc.add_heading("MessageMesh", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Admin End-to-End Flow Documentation")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = ACCENT

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
m = meta.add_run("Administration & Moderation  •  Application URL: http://localhost:5173/")
m.font.size = Pt(10)
m.font.color.rgb = MUTED

doc.add_paragraph()

intro = doc.add_paragraph()
intro.add_run(
    "This document walks through the complete administrator journey of the MessageMesh "
    "real-time chat application. Starting from an admin sign-in, it covers the admin "
    "navigation, the dashboard with its twelve live platform-analytics tiles, and full "
    "user management — searching and filtering by role and status, adjustable page sizes, "
    "creating accounts, viewing a user's details, forcing logout, resetting passwords, "
    "promoting/demoting, and bulk activate/deactivate/delete. It then covers conversation "
    "oversight and moderation (filtering, message-level moderation, deletion and restore), "
    "the audit log with date-range filtering and CSV export, and theme switching. Each "
    "step includes a screenshot captured from the running application and an explanation of "
    "the screen and the underlying functionality. All admin routes are protected "
    "server-side and client-side so that only accounts with the ADMIN role can reach them."
)

doc.add_page_break()

# ---- Steps ----
for i, step in enumerate(steps, start=1):
    h = doc.add_heading(f"Step {i}: {step['title']}", level=1)
    for r in h.runs:
        r.font.color.rgb = ACCENT

    u = doc.add_paragraph()
    ur = u.add_run(f"Route: {step['url']}")
    ur.italic = True
    ur.font.size = Pt(9)
    ur.font.color.rgb = MUTED

    img_path = os.path.join(SHOTS, step["img"])
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(img_path, width=Inches(6.0))

    doc.add_paragraph(step["desc"])
    if i != len(steps):
        doc.add_page_break()

doc.save(OUT)
print("Saved:", OUT)
