"""Generate a Word document documenting the MessageMesh end-to-end UI flow."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(BASE, "screenshots")
OUT = os.path.join(BASE, "MessageMesh-UI-Flow.docx")

ACCENT = RGBColor(0x1F, 0x6F, 0xEB)
MUTED = RGBColor(0x55, 0x5F, 0x6D)

steps = [
    {
        "img": "01-welcome.png",
        "title": "Welcome / Landing Page",
        "url": "http://localhost:5173/",
        "desc": "The public landing page is the entry point of MessageMesh. It presents the "
                "product value proposition (\"See how MessageMesh works\"), highlights key "
                "features such as instant delivery and secure-by-design messaging, and offers "
                "two primary calls to action: \"Get Started\" (registration) and \"Sign in\". "
                "A color-mode toggle and marketing navigation (Product, Security, Pricing, "
                "Resources) are available in the header.",
    },
    {
        "img": "02-login.png",
        "title": "Sign In Page",
        "url": "http://localhost:5173/login",
        "desc": "Clicking \"Sign in\" routes the user to the login screen. It provides a clean "
                "form with Username and Password fields, a primary \"Sign in\" button, and a "
                "link to create a new account. Authentication is JWT-based; on success a token "
                "is stored and used for all subsequent API and WebSocket requests.",
    },
    {
        "img": "03-login-filled.png",
        "title": "Entering Credentials",
        "url": "http://localhost:5173/login",
        "desc": "The user enters their credentials — username \"ram\" and the account "
                "password. The password field is masked for security. The \"Sign in\" button "
                "submits the credentials to the backend authentication endpoint, which returns a "
                "JWT for the standard (non-admin) user account.",
    },
    {
        "img": "04-chats-list.png",
        "title": "Chats Home / Conversation List",
        "url": "http://localhost:5173/chats",
        "desc": "After a successful login the signed-in user (\"Ram H\") lands on the main chat "
                "workspace. The left sidebar lists the user's conversations with avatars, the "
                "latest message preview and a timestamp — here a direct message with \"Admin "
                "User\". A search box allows filtering conversations, and the \"New\" button starts "
                "a new chat. The right pane shows an empty state prompting the user to select or "
                "start a conversation.",
    },
    {
        "img": "05-conversation.png",
        "title": "Opening a Conversation",
        "url": "http://localhost:5173/chats/{id}",
        "desc": "Selecting a conversation (here the direct message with \"Admin User\") opens "
                "the message panel. The header shows the conversation name and type (Direct "
                "message), along with actions to search messages and open conversation options. "
                "Messages are grouped by day (\"Today\"), display the time, and the message "
                "composer is docked at the bottom.",
    },
    {
        "img": "06-compose-message.png",
        "title": "Composing a Message",
        "url": "http://localhost:5173/chats/{id}",
        "desc": "The user types a message into the composer. The input supports Enter to send and "
                "Shift+Enter for a new line. Once text is entered, the \"Send\" button becomes "
                "enabled.",
    },
    {
        "img": "07-message-sent.png",
        "title": "Message Sent (Real-Time Delivery)",
        "url": "http://localhost:5173/chats/{id}",
        "desc": "After sending, the message appears immediately in the thread aligned to the right "
                "(the current user's messages). A delivery status indicator transitions from "
                "\"Sending\" to \"Sent\". The message is delivered in real time over the WebSocket "
                "(STOMP) connection and the conversation's preview and timestamp in the sidebar "
                "update accordingly.",
    },
    {
        "img": "08-conversation-options.png",
        "title": "Conversation Options",
        "url": "http://localhost:5173/chats/{id}",
        "desc": "Opening the conversation options menu (the \"⋯\" control in the conversation "
                "header) reveals per-conversation actions available to the user: Mute to silence "
                "notifications, Archive to move the chat out of the active list, and Delete "
                "conversation to remove it. This lets the user manage and declutter their own "
                "conversations.",
    },
    {
        "img": "09-new-conversation.png",
        "title": "New Conversation Modal",
        "url": "http://localhost:5173/chats",
        "desc": "The \"New\" button opens the New Conversation modal. Users can select one or more "
                "people (via checkboxes) to start a one-to-one chat or create a group. Each "
                "contact is shown with an avatar, display name and handle.",
    },
    {
        "img": "10-dark-mode.png",
        "title": "Dark Mode",
        "url": "http://localhost:5173/chats/{id}",
        "desc": "The color-mode toggle switches the entire application between light and dark "
                "themes. The dark theme re-styles the header, sidebar, message bubbles and "
                "composer while preserving layout and readability — useful for low-light "
                "environments.",
    },
    {
        "img": "11-profile-menu.png",
        "title": "User Profile Menu / Sign Out",
        "url": "http://localhost:5173/chats/{id}",
        "desc": "Clicking the user avatar in the top-right opens the profile menu, showing the "
                "signed-in user's name and handle along with a \"Sign out\" action. Signing out "
                "clears the JWT session and returns the user to the public landing/login flow, "
                "completing the end-to-end journey.",
    },
]

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

# ---- Title page ----
title = doc.add_heading("MessageMesh", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("End-to-End UI Flow Documentation")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = ACCENT

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
m = meta.add_run("Real-Time Chat Application  •  Application URL: http://localhost:5173/")
m.font.size = Pt(10)
m.font.color.rgb = MUTED

doc.add_paragraph()

intro = doc.add_paragraph()
intro.add_run(
    "This document walks through the complete user journey of the MessageMesh real-time "
    "chat application, from the public landing page through authentication, browsing "
    "conversations, sending messages in real time, viewing group details, starting new "
    "conversations, switching themes and finally signing out. Each step includes a "
    "screenshot captured from the running application and an explanation of the screen."
)

doc.add_page_break()

# ---- Steps ----
for i, step in enumerate(steps, start=1):
    h = doc.add_heading(f"Step {i}: {step['title']}", level=1)
    for r in h.runs:
        r.font.color.rgb = ACCENT

    u = doc.add_paragraph()
    ur = u.add_run(f"Route: {step['url']}")
    ur.italic = True
    ur.font.size = Pt(9)
    ur.font.color.rgb = MUTED

    img_path = os.path.join(SHOTS, step["img"])
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(img_path, width=Inches(6.0))

    doc.add_paragraph(step["desc"])
    if i != len(steps):
        doc.add_page_break()

doc.save(OUT)
print("Saved:", OUT)
